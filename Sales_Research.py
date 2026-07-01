!pip install python-docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Document banao
doc = Document()

# ============ STYLES =============
def add_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x4A, 0x25, 0x11)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(0x6F, 0x4E, 0x37)

def add_para(doc, text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(10)
    para_format = para.paragraph_format
    para_format.line_spacing = Pt(18)

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_numbered(doc, text):
    doc.add_paragraph(text, style='List Number')

# ============ TITLE PAGE =============
doc.add_paragraph()
doc.add_paragraph()
add_title(doc, "Sales Trend and Time-Based Performance Analysis")
add_title(doc, "for Afficionado Coffee Roasters")
doc.add_paragraph()
para = doc.add_paragraph("A Data Analytics Research Paper")
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()
para2 = doc.add_paragraph("Prepared as part of the Unified Mentor Internship Program")
para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
para3 = doc.add_paragraph("Author: Yashika")
para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
para3.runs[0].bold = True
para4 = doc.add_paragraph("2026")
para4.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============ ABSTRACT =============
add_heading(doc, "Abstract", level=1)
add_para(doc, "Afficionado Coffee Roasters operates three retail locations in New York City — Lower Manhattan, Hell's Kitchen, and Astoria — and generates large volumes of point-of-sale transaction data. This research paper presents a comprehensive time-based performance analysis of 149,116 transactions recorded in 2025. Using Python, Excel, Tableau, Power BI, and Streamlit, the analysis identifies clear temporal demand patterns, quantifies revenue concentration during morning hours, and surfaces actionable recommendations for staffing and operational planning.")
add_para(doc, "Key findings show that the morning period (6 AM–11 AM) generates 55.6% of total revenue, transaction volume peaks between 8 AM and 10 AM, and Hell's Kitchen is the highest-revenue location.")

# ============ 1. INTRODUCTION =============
add_heading(doc, "1. Introduction", level=1)
add_heading(doc, "1.1 Background", level=2)
add_para(doc, "In specialty coffee retail, when sales occur is just as important as what is sold. Poor understanding of temporal demand patterns can result in overstaffing during slow hours, understaffing during rush periods, inconsistent customer experience, and inefficient operational costs.")

add_heading(doc, "1.2 Problem Statement", level=2)
add_para(doc, "Despite having transaction-level data, Afficionado Coffee Roasters currently lacks:")
add_bullet(doc, "A consolidated view of sales trends over time")
add_bullet(doc, "Clear identification of peak and off-peak hours")
add_bullet(doc, "Hourly demand insights across different store locations")
add_para(doc, "Without structured analytics, operational decisions are reactive and experience-driven rather than data-driven.")

add_heading(doc, "1.3 Objectives", level=2)
add_para(doc, "Primary Objectives:")
add_bullet(doc, "Identify overall sales trends across 2025")
add_bullet(doc, "Determine the busiest and slowest transaction hours")
add_bullet(doc, "Identify peak transaction hours across all locations")
add_para(doc, "Secondary Objectives:")
add_bullet(doc, "Compare revenue performance across store locations")
add_bullet(doc, "Quantify the revenue contribution of each time-of-day segment")
add_bullet(doc, "Build interactive dashboards for ongoing stakeholder use")

doc.add_page_break()

# ============ 2. DATASET DESCRIPTION =============
add_heading(doc, "2. Dataset Description", level=1)
add_para(doc, "The dataset consists of 149,116 individual transaction records collected from three Afficionado Coffee Roasters locations throughout 2025. The data was clean on arrival, with zero missing values across all 11 original columns.")

# Table - Dataset fields
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Field'
hdr_cells[1].text = 'Description'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('transaction_id', 'Unique identifier for each transaction'),
    ('transaction_time', 'Time of day the transaction occurred (HH:MM:SS)'),
    ('transaction_qty', 'Number of units purchased'),
    ('store_location', 'Store: Lower Manhattan, Hells Kitchen, or Astoria'),
    ('unit_price', 'Price per unit (USD)'),
    ('product_category', 'High-level product grouping (Coffee, Tea, Bakery etc.)'),
    ('product_type', 'Granular product classification'),
]
for field, desc in data:
    row_cells = table.add_row().cells
    row_cells[0].text = field
    row_cells[1].text = desc

doc.add_paragraph()

# ============ 3. METHODOLOGY =============
add_heading(doc, "3. Methodology", level=1)
add_heading(doc, "3.1 Data Ingestion and Validation", level=2)
add_para(doc, "The dataset was loaded using Python (Pandas) and validated for timestamp format consistency, duplicate transaction IDs, and logical consistency. No data quality issues were found.")

add_heading(doc, "3.2 Feature Engineering", level=2)
add_para(doc, "The following derived features were created:")
add_bullet(doc, "Revenue = transaction_qty x unit_price")
add_bullet(doc, "Hour of day (0-23), extracted from transaction_time")
add_bullet(doc, "Time buckets: Morning (6-11), Afternoon (12-16), Evening (17-21), Late Hours (22-5)")

add_heading(doc, "3.3 Analytical Tools", level=2)
add_bullet(doc, "Python (Pandas, Matplotlib, Seaborn) - data cleaning, EDA, statistical summaries")
add_bullet(doc, "Microsoft Excel - pivot tables and PivotCharts")
add_bullet(doc, "Tableau - interactive heatmaps and published public dashboard")
add_bullet(doc, "Power BI - executive-level KPI dashboard")
add_bullet(doc, "Streamlit - live, filterable web application")

doc.add_page_break()

# ============ 4. FINDINGS =============
add_heading(doc, "4. Analysis and Findings", level=1)

add_heading(doc, "4.1 Overall Performance Summary", level=2)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Metric'
hdr2[1].text = 'Value'
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True

summary_data = [
    ('Total Transactions', '149,116'),
    ('Total Revenue', '$698,812.33'),
    ('Average Revenue per Transaction', '$4.69'),
    ('Peak Transaction Hour', '10:00 AM (all stores)'),
    ('Top Revenue Store', "Hell's Kitchen ($236,511)"),
    ('Top Time Bucket', 'Morning 6-11 AM ($388,288)')
]
for metric, value in summary_data:
    row = table2.add_row().cells
    row[0].text = metric
    row[1].text = value

doc.add_paragraph()

add_heading(doc, "4.2 Store-Level Performance", level=2)
table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(['Store', 'Revenue', 'Transactions', 'Avg/Txn']):
    hdr3[i].text = h
    hdr3[i].paragraphs[0].runs[0].bold = True

store_data = [
    ("Hell's Kitchen", "$236,511", "50,735", "$4.66"),
    ("Astoria", "$232,243", "50,599", "$4.59"),
    ("Lower Manhattan", "$230,057", "47,782", "$4.81")
]
for row_data in store_data:
    row = table3.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
add_para(doc, "All three stores operate within 3% of each other in total revenue, indicating balanced performance across the chain.")

add_heading(doc, "4.3 Time-of-Day Analysis", level=2)
table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
for i, h in enumerate(['Time Bucket', 'Revenue', '% of Total']):
    hdr4[i].text = h
    hdr4[i].paragraphs[0].runs[0].bold = True

time_data = [
    ("Morning (6-11)", "$388,288", "55.6%"),
    ("Afternoon (12-16)", "$204,720", "29.3%"),
    ("Evening (17-21)", "$105,802", "15.1%")
]
for row_data in time_data:
    row = table4.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
add_para(doc, "Morning alone generates more than half of all daily revenue. Transaction volume peaks sharply at 8-10 AM, then drops 45% after 11 AM.")

add_heading(doc, "4.4 Cross-Location Heatmap Analysis", level=2)
add_para(doc, "Hourly heatmap analysis confirms the 8-10 AM peak is consistent across all three locations, indicating a shared commuter and pre-work coffee purchase pattern. Hell's Kitchen shows the most pronounced peak-to-trough swing.")

doc.add_page_break()

# ============ 5. RECOMMENDATIONS =============
add_heading(doc, "5. Recommendations", level=1)
add_numbered(doc, "Concentrate staffing during 7:30 AM-10:30 AM across all stores to match the morning demand surge.")
add_numbered(doc, "Reduce staffing levels after 7 PM when transaction volume drops to very low levels.")
add_numbered(doc, "Consider targeted afternoon promotions (12-4 PM) to lift the underutilized revenue window.")
add_numbered(doc, "Replicate Hell's Kitchen operational practices at other locations given its leading performance.")
add_numbered(doc, "Adopt the Streamlit dashboard for self-service monitoring so management can track patterns without manual analysis.")

# ============ 6. CONCLUSION =============
add_heading(doc, "6. Conclusion", level=1)
add_para(doc, "This project provides a clear, evidence-based understanding of when sales occur at Afficionado Coffee Roasters. By uncovering temporal demand patterns across hours, time-of-day segments, and store locations, it equips management with actionable insights to improve staffing efficiency, operational planning, and overall retail performance.")
add_para(doc, "The combination of Python, Excel, Tableau, Power BI, and Streamlit demonstrates how the same dataset can be analyzed through complementary lenses — from deep statistical exploration to live, interactive business dashboards — supporting both one-time research and ongoing operational decision-making.")

# ============ 7. DELIVERABLES =============
add_heading(doc, "7. Deliverables", level=1)
add_bullet(doc, "Research Paper - this document (EDA, insights, and recommendations)")
add_bullet(doc, "Streamlit Dashboard - live, interactive analytics application")
add_bullet(doc, "Tableau Public Dashboard - published interactive visualization")
add_bullet(doc, "Power BI Dashboard - executive summary view")
add_bullet(doc, "GitHub Repository - full source code, cleaned dataset, and documentation")

# ============ SAVE =============
doc.save('Afficionado_Research_Paper.docx')
print("Research paper saved successfully!")
print("Download: Afficionado_Research_Paper.docx")